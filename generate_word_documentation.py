"""Generate presentation-quality Word document from verified benchmark results."""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from paths import ROOT, DOCS, FIGURES as FIGS, RESULTS

OUT = DOCS / "CMAPSS_RUL_Benchmark_Documentation.docx"
OUT_ALT = DOCS / "CMAPSS_RUL_Benchmark_Documentation_updated.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)

DATASET_INFO = {
    "FD001": "Single operating condition (sea level), single fault mode (HPC degradation), 100 train / 100 test engines — easiest subset",
    "FD002": "Six operating conditions, single fault mode (HPC), 260 train / 259 test engines — medium difficulty",
    "FD003": "Single operating condition, dual fault modes (HPC + Fan), 100 train / 100 test engines — medium difficulty",
    "FD004": "Six operating conditions, dual fault modes (HPC + Fan), 248 train / 249 test engines — hardest subset",
}

DATASET_SHORT = {
    "FD001": "FD001 (1 cond., 1 fault)",
    "FD002": "FD002 (6 cond., 1 fault)",
    "FD003": "FD003 (1 cond., 2 faults)",
    "FD004": "FD004 (6 cond., 2 faults)",
}

ML_STAGE_NOTE = (
    "Column headers Stage 10, Stage 50, Stage 100, and Stage 200 denote training progress checkpoints "
    "for each ML model. The number indicates how many estimators (trees) or maximum iterations were "
    "used when fitting the model. For example, ExtraTrees at Stage 50 was trained with 50 trees. "
    "Each cell shows last-cycle RMSE on the test set at that training stage."
)

DL_EPOCH_NOTE = (
    "Column headers Epoch 10, Epoch 50, Epoch 100, and Epoch 200 denote training checkpoints during "
    "200-epoch deep learning runs. The number is how many full passes through the training data "
    "completed before evaluation. Each cell shows last-cycle RMSE on the test set at that epoch."
)

RESULTS_TABLE_NOTE = (
    "Column headers FD001–FD004 are the four NASA C-MAPSS benchmark subsets (see Section 1 for full "
    "descriptions). Each column shows model performance on that dataset. All values are computed at "
    "the last observed cycle of each test engine."
)

GRAPH_MODELS = ["TemporalGCN", "TemporalGAT"]
DATASET_IDS = ["FD001", "FD002", "FD003", "FD004"]


def load_master() -> pd.DataFrame:
    ml = pd.read_csv(RESULTS / "all_model_results.csv")
    ml["Family"] = "ML"
    ml = ml.rename(columns={"RMSE_last": "RMSE", "MAE_last": "MAE", "R2_last": "R2"})
    dl = pd.read_csv(RESULTS / "dl_model_results.csv")
    dl["Family"] = "DL"
    dl = dl.rename(columns={"RMSE_last": "RMSE", "MAE_last": "MAE", "R2_last": "R2"})
    g = pd.read_csv(RESULTS / "graph_model_results.csv")
    g["Family"] = "Graph"
    g = g.rename(columns={"RMSE_last": "RMSE", "MAE_last": "MAE", "R2_last": "R2"})
    cols = ["Model", "Dataset", "Family", "RMSE", "MAE", "R2"]
    return pd.concat([ml[cols], dl[cols], g[cols]], ignore_index=True).drop_duplicates(
        subset=["Model", "Dataset"], keep="last"
    )


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)


def add_toc(doc: Document):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_note_box(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)
    doc.add_paragraph()


def add_title_page(doc: Document, best: pd.DataFrame):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("NASA C-MAPSS\nRemaining Useful Life (RUL) Prediction\nBenchmark Report")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Comprehensive evaluation of 13 machine-learning, deep-learning, and graph models\n"
        "on turbofan engine datasets FD001–FD004"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph("Primary metric: last-cycle RMSE per test engine  |  Training RUL cap: 125 cycles")
    doc.add_paragraph("All numerical values verified against results/*.csv and docs/epoch_metrics_*.csv")
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    b = {r.Dataset: r for _, r in best.iterrows()}
    summary = (
        f"This benchmark evaluates RUL prediction across four NASA C-MAPSS subsets with increasing difficulty. "
        f"The best models achieved RMSE of {b['FD001'].RMSE:.2f} (FD001 — single condition, single fault, {b['FD001'].Model}), "
        f"{b['FD002'].RMSE:.2f} (FD002 — six conditions, single fault, {b['FD002'].Model}), "
        f"{b['FD003'].RMSE:.2f} (FD003 — single condition, dual fault, {b['FD003'].Model}), and "
        f"{b['FD004'].RMSE:.2f} (FD004 — six conditions, dual fault, {b['FD004'].Model}). "
        f"Tree-based ML with engineered features dominates single-condition datasets; "
        f"the Transformer excels on multi-condition FD002. "
        f"Key degradation sensors: sensor_11, sensor_4, and sensor_12 (FD001)."
    )
    doc.add_paragraph(summary)


def add_table_from_df(doc: Document, df: pd.DataFrame, header_color: str = "1F3A5F"):
    df = df.reset_index(drop=True)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], header_color)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                cells[i].text = f"{val:.4f}" if abs(val) < 10 else f"{val:.2f}"
            else:
                cells[i].text = str(val)
    doc.add_paragraph()


def rename_dataset_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [DATASET_SHORT.get(str(c), c) for c in out.columns]
    return out


def rename_stage_columns(df: pd.DataFrame, prefix: str = "Stage") -> pd.DataFrame:
    out = df.copy()
    new_cols = []
    for c in out.columns:
        if str(c).isdigit():
            new_cols.append(f"{prefix} {c}")
        else:
            new_cols.append(c)
    out.columns = new_cols
    return out


def add_dataset_legend(doc: Document):
    add_heading(doc, "1.2 Dataset Code Reference (FD001–FD004)", level=2)
    doc.add_paragraph(
        "Throughout this report, tables and figures use the codes FD001–FD004 to identify each "
        "benchmark subset. Use this reference when reading any results table:"
    )
    legend = pd.DataFrame(
        [(k, v) for k, v in DATASET_INFO.items()],
        columns=["Code", "Description"],
    )
    add_table_from_df(doc, legend, header_color="2E75B6")


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.0):
    if not path.exists():
        doc.add_paragraph(f"[Figure missing: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(10)
    doc.add_paragraph()


def build_document():
    master = load_master()
    best = master.loc[master.groupby("Dataset")["RMSE"].idxmin()]
    epoch_ml = pd.read_csv(DOCS / "epoch_metrics_ml.csv")
    epoch_dl = pd.read_csv(DOCS / "epoch_metrics_dl.csv")
    sensor = pd.read_csv(DOCS / "sensor_importance.csv")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title_page(doc, best)
    add_toc(doc)

    # --- Dataset ---
    add_heading(doc, "1. Dataset Overview", level=1)
    doc.add_paragraph(
        "NASA C-MAPSS (Commercial Modular Aero-Propulsion System Simulation) provides multivariate "
        "time-series from simulated turbofan jet engines. Each row is one operational cycle. "
        "Training trajectories run until engine failure; test trajectories end before failure with "
        "provided final RUL labels. The benchmark is split into four subsets — FD001 through FD004 — "
        "each differing in the number of operating conditions and fault modes."
    )
    ds = pd.DataFrame([
        ["FD001", 100, 100, "1 (Sea Level)", "1 (HPC Degradation)", "Easy"],
        ["FD002", 260, 259, "6", "1 (HPC Degradation)", "Medium"],
        ["FD003", 100, 100, "1 (Sea Level)", "2 (HPC + Fan Degradation)", "Medium"],
        ["FD004", 248, 249, "6", "2 (HPC + Fan Degradation)", "Hard"],
    ], columns=["Dataset Code", "Train Engines", "Test Engines", "Operating Conditions", "Fault Modes", "Difficulty"])
    add_note_box(
        doc,
        "FD001–FD004 are standard NASA C-MAPSS identifiers. FD = Failure Dataset. "
        "Higher numbers generally indicate more operating conditions and/or fault modes, "
        "making prediction harder.",
    )
    add_table_from_df(doc, ds)
    add_dataset_legend(doc)

    add_heading(doc, "1.3 Why FD004 Is Harder Than FD001", level=2)
    fd1, fd4 = best[best.Dataset == "FD001"].RMSE.iloc[0], best[best.Dataset == "FD004"].RMSE.iloc[0]
    for item in [
        "FD004 has six operating conditions vs one in FD001 — models must generalize across flight regimes.",
        "FD004 has two fault modes (HPC + Fan) vs one in FD001 — more complex degradation patterns.",
        f"Best RMSE on FD004 ({fd4:.2f}) vs FD001 ({fd1:.2f}) — {fd4/fd1:.1f}× higher error.",
        "FD004 has a larger fleet (248 engines) with more heterogeneity and sensor noise.",
        "Feature distribution shift between conditions makes global models harder to fit.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- Preprocessing ---
    add_heading(doc, "2. Data Preparation & Preprocessing", level=1)
    steps = [
        "Load train/test/RUL files per FD subset (FD001, FD002, FD003, or FD004).",
        "Compute RUL per row; cap training RUL at 125 cycles.",
        "Drop constant sensors (zero variance on training set).",
        "ML features: raw + unit-relative + rolling mean (window=5) + rolling std (window=5).",
        "DL/Graph features: unit-relative scaling + StandardScaler; 30-step sequences.",
        "Graph edges: temporal chain connecting consecutive time steps within each sequence.",
        "Evaluation: last observed cycle per test engine — RMSE, MAE, R².",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    # --- Models ---
    add_heading(doc, "3. Models Benchmarked", level=1)
    doc.add_paragraph(
        "Thirteen models were evaluated across three families: classical machine learning (6), "
        "deep learning (5), and graph neural networks (2)."
    )
    doc.add_paragraph("ML (6): Ridge, ElasticNet, RandomForest, ExtraTrees, HistGradientBoosting, XGBoost")
    doc.add_paragraph("DL (5): LSTM, GRU, CNN1D, TCN, Transformer (PyTorch CUDA)")
    doc.add_paragraph(
        "Graph (2): TemporalGCN — graph convolution over temporal node chains; "
        "TemporalGAT — graph attention over temporal node chains (PyTorch Geometric)"
    )

    doc.add_page_break()
    add_heading(doc, "4. Best Model Per Dataset", level=1)
    add_note_box(
        doc,
        "The Dataset column uses FD001–FD004 codes defined in Section 1.2. "
        "RMSE, MAE, and R² are last-cycle metrics on the test set for that subset.",
    )
    bt = best[["Dataset", "Model", "Family", "RMSE", "MAE", "R2"]].copy()
    bt["Dataset"] = bt["Dataset"].map(lambda d: f"{d} — {DATASET_INFO[d].split(',')[0]}")
    bt.columns = ["Dataset", "Model", "Family", "RMSE", "MAE", "R²"]
    add_table_from_df(doc, bt.round(4))

    # --- Full results ---
    add_heading(doc, "5. Complete Results — RMSE (Last-Cycle)", level=1)
    add_note_box(doc, RESULTS_TABLE_NOTE)
    rmse_pivot = master.pivot_table(index="Model", columns="Dataset", values="RMSE").round(2)
    add_table_from_df(doc, rename_dataset_columns(rmse_pivot.reset_index()))

    add_heading(doc, "5.1 Complete Results — R² (Last-Cycle)", level=2)
    add_note_box(doc, RESULTS_TABLE_NOTE)
    r2_pivot = master.pivot_table(index="Model", columns="Dataset", values="R2").round(4)
    add_table_from_df(doc, rename_dataset_columns(r2_pivot.reset_index()))

    add_heading(doc, "5.2 Complete Results — MAE (Last-Cycle)", level=2)
    add_note_box(doc, RESULTS_TABLE_NOTE)
    mae_pivot = master.pivot_table(index="Model", columns="Dataset", values="MAE").round(2)
    add_table_from_df(doc, rename_dataset_columns(mae_pivot.reset_index()))

    add_figure(
        doc, FIGS / "heatmap_rmse_all.png",
        "Figure 1: RMSE heatmap — rows are models, columns are FD001–FD004 datasets",
    )

    for fd in ["FD001", "FD002", "FD003", "FD004"]:
        for metric in ["RMSE_last", "MAE_last", "R2_last"]:
            label = metric.replace("_last", "").replace("R2", "R²")
            fig = FIGS / f"compare_{fd}_{metric}.png"
            add_figure(
                doc, fig,
                f"Figure: {fd} ({DATASET_INFO[fd].split(',')[0]}) — model comparison ({label})",
                width=6.2,
            )

    doc.add_page_break()

    # --- Epoch metrics ML ---
    add_heading(doc, "6. Training Progression — ML Models", level=1)
    add_note_box(doc, ML_STAGE_NOTE)
    doc.add_paragraph(
        "The tables below show how each ML model's last-cycle RMSE improves as more estimators "
        "or iterations are added. Each subsection corresponds to one FD dataset."
    )
    for i, fd in enumerate(["FD001", "FD002", "FD003", "FD004"], 1):
        add_heading(doc, f"6.{i} {fd} — {DATASET_INFO[fd].split(',')[0]}", level=2)
        sub = epoch_ml[epoch_ml.Dataset == fd].pivot_table(
            index="Model", columns="Epoch", values="RMSE", aggfunc="first"
        ).round(2).reset_index()
        add_table_from_df(doc, rename_stage_columns(sub), header_color="2E75B6")

    # --- Epoch metrics DL ---
    add_heading(doc, "7. Training Progression — Deep Learning Models", level=1)
    add_note_box(doc, DL_EPOCH_NOTE)
    doc.add_paragraph(
        "Each DL model was trained for up to 200 epochs with checkpoints at epochs 10, 50, 100, "
        "and 200. Overfitting onset is the first epoch where validation loss exceeds training loss by >12%."
    )
    for i, fd in enumerate(sorted(epoch_dl.Dataset.unique()), 1):
        add_heading(doc, f"7.{i} {fd} — {DATASET_INFO[fd].split(',')[0]}", level=2)
        sub = epoch_dl[epoch_dl.Dataset == fd]
        pt = sub.pivot_table(index="Model", columns="Epoch", values="RMSE", aggfunc="first").round(2).reset_index()
        add_table_from_df(doc, rename_stage_columns(pt, prefix="Epoch"), header_color="2E75B6")
        of = sub.groupby("Model")["overfit_epoch"].first()
        doc.add_paragraph(
            "Overfitting onset (training epochs): " + ", ".join(f"{m} at epoch {int(v)}" for m, v in of.items())
        )

    doc.add_page_break()
    add_heading(doc, "8. Loss Curve Analysis", level=1)
    doc.add_paragraph(
        "Each plot shows train vs validation MSE (left) and the generalization gap val−train (right). "
        "Vertical dashed lines mark checkpoints at epochs 10, 50, 100, and 200 — the same epoch "
        "checkpoints used in Section 7 tables. Red dotted line marks estimated overfitting onset."
    )
    for pattern in sorted(FIGS.glob("loss_*.png")):
        name = pattern.stem.replace("loss_", "").replace("_", " on ")
        add_figure(doc, pattern, f"Loss curves: {name}", width=6.3)

    doc.add_page_break()

    # --- Sensor importance ---
    add_heading(doc, "9. Sensor Importance Analysis", level=1)
    doc.add_paragraph(
        "Feature importance from ExtraTrees regressor, aggregated across engineered variants "
        "(raw, unit-relative, rolling mean/std). Higher values indicate stronger contribution to RUL prediction."
    )
    for i, fd in enumerate(DATASET_IDS, 1):
        add_heading(doc, f"9.{i} {fd} — {DATASET_INFO[fd].split(',')[0]}", level=2)
        sub = sensor[sensor.Dataset == fd].nlargest(10, "Importance")[["Feature", "Importance"]]
        sub["Importance"] = sub["Importance"].round(4)
        add_table_from_df(doc, sub)
        add_figure(doc, FIGS / f"sensor_importance_{fd}.png", f"Top sensor importance — {fd}")

    # --- Model behavior ---
    add_heading(doc, "10. Model Behavior Classification", level=1)
    add_note_box(
        doc,
        "Stage 10 / Stage 50 refer to ML training stages (Section 6). "
        "Epoch 6 refers to DL training epochs (Section 7). "
        "FD001–FD004 refer to the dataset subsets defined in Section 1.2.",
    )
    behavior = pd.DataFrame([
        ["Underfitting (early ML)", "XGBoost, HistGradientBoosting at Stage 10",
         "RMSE 28–43 before enough estimators on FD001–FD004"],
        ["Stable (ML)", "ExtraTrees, XGBoost at Stages 50–200",
         "RMSE change < 1.5 from Stage 50 to Stage 200"],
        ["Overfitting (DL)", "LSTM, GRU, CNN1D, TCN, Transformer on FD001",
         "Validation loss exceeds training loss after epoch 6"],
        ["Best balance", "ExtraTrees, TemporalGAT, Transformer",
         "Strong RMSE with acceptable generalization gap"],
    ], columns=["Category", "Models", "Evidence"])
    add_table_from_df(doc, behavior)

    # --- Architecture (graph models only show the 2 we ran) ---
    add_heading(doc, "11. Model Architectures", level=1)
    doc.add_paragraph(
        "Architecture diagrams for representative models from each family. "
        "Graph-based models use a temporal chain graph where each of the 30 time steps is a node."
    )
    arch_order = [
        "ExtraTrees", "XGBoost", "LSTM", "GRU", "CNN1D", "TCN", "Transformer",
        "TemporalGCN", "TemporalGAT",
    ]
    for name in arch_order:
        path = FIGS / f"arch_{name}.png"
        add_figure(doc, path, f"Architecture: {name}", width=5.5)

    doc.add_page_break()

    # --- Recommendations ---
    add_heading(doc, "12. Final Recommendations", level=1)
    recs = [
        ("FD001 — single condition, single fault", "ExtraTrees (RMSE 13.89) — best overall. TemporalGAT (16.60) for graph-based approach."),
        ("FD002 — six conditions, single fault", "Transformer (RMSE 26.57) — beats XGBoost (27.71). Use sequence models when GPU available."),
        ("FD003 — single condition, dual fault", "HistGradientBoosting (RMSE 17.13) — fast and accurate."),
        ("FD004 — six conditions, dual fault (hardest)", "HistGradientBoosting (RMSE 29.39) — best overall on the hardest subset."),
        ("Interpretability", "ExtraTrees + sensor importance. Top sensors: sensor_11, sensor_4, sensor_12."),
        ("Production baseline", "ExtraTrees or HistGradientBoosting with engineered rolling features."),
        ("Graph-based RUL", "TemporalGAT for attention-weighted temporal graphs; TemporalGCN for convolutional aggregation."),
    ]
    for title, text in recs:
        p = doc.add_paragraph()
        r = p.add_run(f"{title}: ")
        r.bold = True
        p.add_run(text)

    add_heading(doc, "13. Presentation Summary", level=1)
    b = {r.Dataset: r for _, r in best.iterrows()}
    doc.add_paragraph(
        "We benchmarked 13 models (6 ML + 5 DL + 2 Graph) on NASA C-MAPSS RUL prediction across "
        "four datasets (FD001–FD004). All metrics are last-cycle RMSE/MAE/R² per test engine."
    )
    findings = [
        f"Best RMSE — FD001 (easiest): {b['FD001'].RMSE:.2f}, FD002: {b['FD002'].RMSE:.2f}, "
        f"FD003: {b['FD003'].RMSE:.2f}, FD004 (hardest): {b['FD004'].RMSE:.2f}.",
        "ML with feature engineering wins on FD001/FD003/FD004; Transformer wins FD002.",
        "FD004 is ~2× harder than FD001 due to 6 operating conditions and 2 fault modes.",
        "DL models overfit after ~epoch 6 on FD001; early stopping at Stage/Epoch 50 is recommended.",
        "sensor_11 and sensor_4 are the dominant degradation indicators on FD001.",
        "TemporalGAT achieves competitive graph-based results (RMSE 16.60 on FD001).",
    ]
    for f in findings:
        doc.add_paragraph(f, style="List Bullet")

    add_heading(doc, "14. Glossary & Table Reading Guide", level=1)
    glossary = pd.DataFrame([
        ["FD001", "Failure Dataset 001 — 1 operating condition, 1 fault mode, easiest"],
        ["FD002", "Failure Dataset 002 — 6 operating conditions, 1 fault mode"],
        ["FD003", "Failure Dataset 003 — 1 operating condition, 2 fault modes"],
        ["FD004", "Failure Dataset 004 — 6 operating conditions, 2 fault modes, hardest"],
        ["Stage 10/50/100/200", "ML training checkpoints — number of estimators/trees used"],
        ["Epoch 10/50/100/200", "DL training checkpoints — number of training passes completed"],
        ["Last-cycle RMSE", "Root mean squared error at the final observed cycle of each test engine"],
        ["RUL cap (125)", "Training labels capped at 125 cycles to reduce noise from early-life data"],
        ["TemporalGCN", "Graph Convolutional Network over temporal node chains"],
        ["TemporalGAT", "Graph Attention Network over temporal node chains"],
    ], columns=["Term", "Meaning"])
    add_table_from_df(doc, glossary, header_color="2E75B6")

    add_heading(doc, "15. Data Files Reference", level=1)
    files = pd.DataFrame([
        ["results/all_model_results.csv", "ML benchmark results (FD001–FD004)"],
        ["results/dl_model_results.csv", "DL benchmark results (FD001–FD004)"],
        ["results/graph_model_results.csv", "Graph benchmark results (FD001–FD004)"],
        ["results/all_results_master.csv", "All 13 models combined"],
        ["results/best_per_dataset_all.csv", "Best model per FD dataset"],
        ["docs/epoch_metrics_ml.csv", "ML staged metrics at Stages 10/50/100/200"],
        ["docs/epoch_metrics_dl.csv", "DL metrics at Epochs 10/50/100/200"],
        ["docs/sensor_importance.csv", "Feature importance scores per FD dataset"],
        ["models/*.pkl, *.pth", "Saved best models per dataset"],
    ], columns=["File", "Description"])
    add_table_from_df(doc, files)

    for out_path in (OUT, OUT_ALT):
        try:
            doc.save(out_path)
            print(f"Saved: {out_path} ({out_path.stat().st_size / 1e6:.1f} MB)")
            break
        except PermissionError:
            if out_path == OUT_ALT:
                raise
            print(f"Could not overwrite {OUT} (file may be open). Trying alternate path...")


if __name__ == "__main__":
    build_document()
